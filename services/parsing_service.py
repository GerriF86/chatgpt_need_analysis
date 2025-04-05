import os
import io
from pathlib import Path

def extract_text_from_pdf(file_path_or_bytes):
    """
    Extracts text from a PDF file.
    :param file_path_or_bytes: Path to the PDF file or bytes of PDF file.
    :return: Extracted text as a string.
    """
    text = ""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        fitz = None
    if fitz:
        # Use PyMuPDF if available for more reliable text extraction
        try:
            if isinstance(file_path_or_bytes, (str, Path)):
                doc = fitz.open(str(file_path_or_bytes))
            else:
                # If bytes provided (from reading the file)
                doc = fitz.open(stream=file_path_or_bytes, filetype='pdf')
        except Exception as e:
            raise RuntimeError(f"Failed to open PDF: {e}")
        for page in doc:
            try:
                text += page.get_text()
            except Exception as e:
                # If any page fails, skip that page
                print(f"Warning: failed to get text from a PDF page: {e}")
        doc.close()
    else:
        # Fallback to PyPDF2 if PyMuPDF is not installed
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("No PDF extraction library (PyMuPDF or PyPDF2) is installed.")
        try:
            if isinstance(file_path_or_bytes, (str, Path)):
                reader = PdfReader(str(file_path_or_bytes))
            else:
                # If bytes provided
                pdf_stream = io.BytesIO(file_path_or_bytes)
                reader = PdfReader(pdf_stream)
        except Exception as e:
            raise RuntimeError(f"Failed to open PDF with PyPDF2: {e}")
        # Iterate through pages
        for page in reader.pages:
            try:
                text += page.extract_text() or ""
            except Exception as e:
                print(f"Warning: failed to extract text from a PDF page with PyPDF2: {e}")
    return text

def extract_text_from_docx(file_path_or_bytes):
    """
    Extracts text from a DOCX file.
    :param file_path_or_bytes: Path to the docx file or bytes of the docx file.
    :return: Extracted text as a string.
    """
    try:
        import docx  # python-docx
    except ImportError:
        raise ImportError("python-docx is not installed, cannot parse .docx files.")
    try:
        if isinstance(file_path_or_bytes, (str, Path)):
            doc = docx.Document(file_path_or_bytes)
        else:
            doc = docx.Document(io.BytesIO(file_path_or_bytes))
    except Exception as e:
        raise RuntimeError(f"Failed to open DOCX file: {e}")
    full_text = []
    for para in doc.paragraphs:
        if para.text:
            full_text.append(para.text)
    return "\n".join(full_text)

def extract_text_from_txt(file_path_or_bytes, encoding="utf-8"):
    """
    Reads text from a TXT file.
    :param file_path_or_bytes: Path to the text file or bytes of text file.
    :param encoding: Encoding to use for decoding bytes. Defaults to 'utf-8'.
    :return: Text content as a string.
    """
    try:
        if isinstance(file_path_or_bytes, (str, Path)):
            with open(file_path_or_bytes, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
        else:
            text = file_path_or_bytes.decode(encoding, errors='ignore')
    except Exception as e:
        raise RuntimeError(f"Failed to read text file: {e}")
    return text

def parse_file(file, file_name=None):
    """
    Determine file type and extract text accordingly.
    :param file: File path, bytes, or file-like object.
    :param file_name: Optional file name (for bytes or file-like) to help determine file type.
    :return: Extracted text as a string.
    """
    # If file is a file-like object (e.g., UploadedFile), read content and recurse with bytes
    if hasattr(file, 'read') and not isinstance(file, (str, Path, bytes, bytearray)):
        try:
            content = file.read()
        except Exception as e:
            raise RuntimeError(f"Failed to read file-like object: {e}")
        if file_name is None and hasattr(file, 'name'):
            file_name = file.name
        return parse_file(content, file_name)
    # Determine extension
    ext = None
    if file_name:
        ext = Path(file_name).suffix.lower()
    elif isinstance(file, (str, Path)):
        ext = Path(file).suffix.lower()
    # Check content signature if extension is missing
    if not ext:
        head_bytes = b''
        if isinstance(file, (bytes, bytearray)):
            head_bytes = file[:5]
        elif hasattr(file, '__iter__'):
            try:
                iterator = iter(file)
                head_bytes = bytes([next(iterator) for _ in range(5)])
            except Exception:
                head_bytes = b''
        if head_bytes.startswith(b'%PDF'):
            ext = '.pdf'
        else:
            ext = '.txt'
    # Dispatch to appropriate extractor
    if ext == '.pdf':
        return extract_text_from_pdf(file)
    elif ext == '.docx':
        return extract_text_from_docx(file)
    elif ext == '.txt':
        return extract_text_from_txt(file)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
