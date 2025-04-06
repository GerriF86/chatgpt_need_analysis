# services/fileparser.py

import io
from pathlib import Path
from typing import Union
import streamlit as st
import PyPDF2
from sklearn.feature_extraction.text import TfidfVectorizer
import re
from docx import Document
# Simulated session state key categories (trimmed down for demo)
SESSION_KEYS = {
    "company_name": ["company", "about us", "who we are", "our company"],
    "job_title": ["job title", "position", "role", "vacancy"],
    "location": ["location", "workplace", "site"],
    "responsibilities": ["responsibilities", "tasks", "your role", "what you will do"],
    "skills": ["requirements", "skills", "what we expect", "qualifications"],
    "benefits": ["benefits", "what we offer", "perks"],
    "application_process": ["how to apply", "recruitment process", "next steps"],
}

def extract_text_from_pdf(pdf_file):
    """Extract raw text from uploaded PDF."""
    reader = PyPDF2.PdfReader(pdf_file)
    text = " ".join(page.extract_text() or "" for page in reader.pages)
    return text.lower()

def match_and_store_keys(text, session_keys):
    """Match content against session keys and store them."""
    for key, patterns in session_keys.items():
        # Search using pattern hints
        for pattern in patterns:
            match = re.search(rf"{pattern}[:\-]?\s*(.+?)(?=\n|\.|$)", text)
            if match:
                value = match.group(1).strip()
                st.session_state[key] = value
                break  # stop on first match

def analyse_pdf_and_store_keys(pdf_file):
    """Full process: extract text → match patterns → store in session_state."""
    raw_text = extract_text_from_pdf(pdf_file)
    match_and_store_keys(raw_text, SESSION_KEYS)
    return raw_text  # for optional inspection or GPT post-processing


def extract_text_from_docx(file_path_or_bytes: Union[str, bytes]) -> str:
    """
    Extract text from a DOCX file given a file path or bytes.
    """
    try:
        if isinstance(file_path_or_bytes, (str, Path)):
            doc = Document(file_path_or_bytes)
        else:
            doc = Document(io.BytesIO(file_path_or_bytes))
        text = "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX file: {e}")
    return text

def extract_text_from_txt(file_path_or_bytes: Union[str, bytes], encoding: str = "utf-8") -> str:
    """
    Read text from a plain text file given a file path or bytes.
    """
    try:
        if isinstance(file_path_or_bytes, (str, Path)):
            with open(file_path_or_bytes, 'r', encoding=encoding) as f:
                text = f.read()
        else:
            text = file_path_or_bytes.decode(encoding, errors='ignore')
    except Exception as e:
        raise RuntimeError(f"Failed to read text file: {e}")
    return text

def parse_file(file: Union[str, bytes, io.IOBase], file_name: str = None) -> str:
    """
    Determine file type and extract text accordingly.
    :param file: File path, bytes, or a file-like object (e.g., an uploaded file).
    :param file_name: Optional file name to help determine file type if file is not a path.
    :return: Extracted text content as a string.
    """
    # If a file-like object is provided, read its content first
    if hasattr(file, 'read') and not isinstance(file, (str, Path, bytes, bytearray)):
        try:
            content = file.read()
        except Exception as e:
            raise RuntimeError(f"Failed to read file-like object: {e}")
        if file_name is None and hasattr(file, 'name'):
            file_name = file.name
        return parse_file(content, file_name)
    # Determine file extension
    ext = None
    if file_name:
        ext = Path(file_name).suffix.lower()
    elif isinstance(file, (str, Path)):
        ext = Path(file).suffix.lower()
    # If no extension is found and we have bytes, attempt to detect PDF signature
    if not ext:
        head_bytes = b''
        if isinstance(file, (bytes, bytearray)):
            head_bytes = file[:5]
        elif hasattr(file, '__iter__') and not isinstance(file, (str, Path)):
            try:
                iterator = iter(file)
                head_bytes = bytes([next(iterator) for _ in range(5)])
            except Exception:
                head_bytes = b''
        if head_bytes.startswith(b'%PDF'):
            ext = '.pdf'
        else:
            ext = '.txt'
    # Dispatch to appropriate extractor based on extension
    if ext == '.pdf':
        return extract_text_from_pdf(file)
    elif ext == '.docx':
        return extract_text_from_docx(file)
    elif ext == '.txt':
        return extract_text_from_txt(file)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
